import os
import docx


def execute(sourceFilePath, finalFileSavePath, generateFileName, contractType):
    os.chdir(finalFileSavePath)

    contractHtmlFileName = generateFileName + "_html.txt"
    contractTxtFileName = generateFileName + ".txt"
    contractFtlFileName = generateFileName + ".ftl"

    sourceFile = docx.Document(sourceFilePath)

    # 生成电子合同模板文件
    # 若是'wb'就表示写二进制文件
    contractTxtFile = open(contractTxtFileName, 'w', encoding='utf8')
    # 合同展示模板文件
    contractHtmlFile = open(contractHtmlFileName, 'w', encoding='utf8')
    # 发标系统合同展示模板文件
    contractFtlFile = open(contractFtlFileName, 'w', encoding='utf8')

    readFile = False

    isPersonal = False

    if "个人" == contractType:
        isPersonal = True

    contractTxtFile.write("借款协议\n")
    contractTxtFile.write("合同编号: ${contractNum}$\n")
    contractTxtFile.write("\n")

    if isPersonal:
        contractTxtFile.write("甲方（借款人）: ${borrowerName}$\n身份证号：${idcardNo}$\n")
    else:
        contractTxtFile.write("甲方（借款人）: ${borrowerName}$\n法定代表人姓名： ${zqzrr}$\n联系地址： ${borrowerAddress}$\n")

    contractTxtFile.write("\n")
    contractTxtFile.write("乙方（出借人）：${name}$\n身份证号码：${idCard}$\n")
    contractTxtFile.write("\n")
    contractTxtFile.write("丙方：浙江金麦穗互联网金融信息服务有限公司\n法定代表人姓名：曾文杰\n联系地址：浙江省杭州市下城区绍兴路161号野风现代中心北楼5A01室\n")
    contractTxtFile.write("\n")
    # 输出每一段的内容
    for para in sourceFile.paragraphs:
        paraText = para.text

        if "START" in paraText:
            readFile = True
            continue

        if "END" in paraText:
            readFile = False
            continue

        if not readFile:
            continue

        if readFile:
            if "(holder)" in paraText:
                paraText = paraText.replace("(holder)", "${holder}$")
            if "(amount)" in paraText:
                paraText = paraText.replace("(amount)", "${amount}$")
            if "(bigAmount)" in paraText:
                paraText = paraText.replace("(bigAmount)", "${bigAmount}$")
            if "(borrowPurpose)" in paraText:
                paraText = paraText.replace("(borrowPurpose)", "${dyw}$")
            if "(baseRate)" in paraText:
                paraText = paraText.replace("(baseRate)", "${annualRate}$")
            if "(floatRate)" in paraText:
                paraText = paraText.replace("(floatRate)", "${couponRate}$")
            if "(bankAccountName)" in paraText:
                paraText = paraText.replace("(bankAccountName)", "${bankAccName}$")
            if "(bankAccountNo)" in paraText:
                paraText = paraText.replace("(bankAccountNo)", "${bankAccNo}$")
            if "(bankName)" in paraText:
                paraText = paraText.replace("(bankName)", "${bankName}$")

        contractTxtFile.write(paraText.strip()+"\n")

    contractTxtFile.write("甲方： ${borrowerName}$\n日期：${nowDays}$\n\n" +
                          "乙方：（在此盖章）\n日期：${nowDays}$\n\n" +
                          "丙方：浙江金麦穗互联网金融信息服务有限公司\n法定代表人：曾文杰\n日期：${nowDays}$")

    if isPersonal:
        contractHtmlFile.write("<div id=\"content\" style=\"max-width:420px;font-family:'微软雅黑';\">\n" +
                               "<style>\n" +
                               "*{margin:0;padding:0}\n" +
                               "</style>\n" +
                               "<h3 align=\"center\" style=\"color:#727171;padding-top:10px;\">借款协议</h3>\n" +
                               "<div style=\"color:#888888;word-break:break-all;padding:8px;line-height:30px;\"> \n" +
                               "<p align=\"right\">合同编号：${contractNum}$</p>\n" +
                               "<p>甲方（借款人）:${borrowerName}$</p>\n" +
                               "<p>身份证号：${idcardNo}$</p>\n" +
                               "<br/>\n" +
                               "<p>乙方（出借人）：_____</p>\n" +
                               "<p>身份证号码：_____</p>\n" +
                               "<br/>\n" +
                               "<p>丙方：浙江金麦穗互联网金融信息服务有限公司</p>\n" +
                               "<p>法定代表人姓名：曾文杰</p>\n" +
                               "<p>联系地址：浙江省杭州市下城区绍兴路161号野风现代中心北楼5A01室</p>\n" +
                               "<br/>\n")

        contractFtlFile.write("<div id=\"content\" style=\"max-width:420px;font-family:'微软雅黑';\">\n" +
                              "<style>\n" +
                              "*{margin:0;padding:0}\n" +
                              "</style>\n" +
                              "<h3 align=\"center\" style=\"color:#727171;padding-top:10px;\">借款协议</h3>\n" +
                              "<div style=\"color:#888888;word-break:break-all;padding:8px;line-height:30px;\"> \n" +
                              "<p align=\"right\">合同编号：${(blselfitem.htbh)!\"________\"}</p>\n" +
                              "<p>甲方（借款人）:${(blselfitem.realname)!\"________\"}</p>\n" +
                              "<p>身份证号：${(blselfitem.idcardno)!\"________\"}\n</p>\n" +
                              "<br/>\n" +
                              "<p>乙方（出借人）：${(account.realname)!\"________\"}</p>\n" +
                              "<p>身份证号码：${(account.idcardno)!\"________\"}</p>\n" +
                              "<br/>\n" +
                              "<p>丙方：浙江金麦穗互联网金融信息服务有限公司</p>\n" +
                              "<p>法定代表人姓名：曾文杰</p>\n" +
                              "<p>联系地址：浙江省杭州市下城区绍兴路161号野风现代中心北楼5A01室</p>\n" +
                              "<br/>\n")

    if not isPersonal:
        contractHtmlFile.write("<div id=\"content\" style=\"max-width:420px;font-family:'微软雅黑';\">\n" +
                               "<style>\n" +
                               "*{margin:0;padding:0}\n" +
                               "</style>\n" +
                               "<h3 align=\"center\" style=\"color:#727171;padding-top:10px;\">借款协议</h3>\n" +
                               "<div style=\"color:#888888;word-break:break-all;padding:8px;line-height:30px;\"> \n" +
                               "<p align=\"right\">合同编号：${contractNum}$</p>\n" +
                               "<p>甲方（借款人）:${borrowerName}$</p>\n" +
                               "<p>法定代表人姓名：${zqzrr}$</p>\n" +
                               "<p>联系地址：${borrowerAddress}$</p>\n" +
                               "<br/>\n" +
                               "<p>乙方（出借人）：_____</p>\n" +
                               "<p>身份证号码：_____</p>\n" +
                               "<br/>\n" +
                               "<p>丙方：浙江金麦穗互联网金融信息服务有限公司</p>\n" +
                               "<p>法定代表人姓名：曾文杰</p>\n" +
                               "<p>联系地址：浙江省杭州市下城区绍兴路161号野风现代中心北楼5A01室</p>\n" +
                               "<br/>\n")

        contractFtlFile.write("<div id=\"content\" style=\"max-width:420px;font-family:'微软雅黑';\">\n" +
                              "<style>\n" +
                              "*{margin:0;padding:0}\n" +
                              "</style>\n" +
                              "<h3 align=\"center\" style=\"color:#727171;padding-top:10px;\">借款协议</h3>\n" +
                              "<div style=\"color:#888888;word-break:break-all;padding:8px;line-height:30px;\"> \n" +
                              "<p align=\"right\">合同编号：${(blselfitem.htbh)!\"________\"}</p>\n" +
                              "<p>甲方（借款人）:${(blselfitem.realname)!\"________\"}</p>\n" +
                              "<p>法定代表人姓名：${(blselfitem.zqzrr)!\"________\"}</p>\n" +
                              "<p>联系地址：${(blselfitem.address)!\"________\"}</p>\n" +
                              "<br/>\n" +
                              "<p>乙方（出借人）：${(account.realname)!\"________\"}</p>\n" +
                              "<p>身份证号码：${(account.idcardno)!\"________\"}</p>\n" +
                              "<br/>\n" +
                              "<p>丙方：浙江金麦穗互联网金融信息服务有限公司</p>\n" +
                              "<p>法定代表人姓名：曾文杰</p>\n" +
                              "<p>联系地址：浙江省杭州市下城区绍兴路161号野风现代中心北楼5A01室</p>\n" +
                              "<br/>\n")


    # 输出每一段的内容
    for para in sourceFile.paragraphs:
        paraText = para.text

        if "START" in paraText:
            readFile = True
            continue

        if "END" in paraText:
            readFile = False
            continue

        if not readFile:
            continue

        if readFile:
            if "(holder)" in paraText:
                paraText = paraText.replace("(holder)", "${holder}$")
            if "(amount)" in paraText:
                paraText = paraText.replace("(amount)", "______")
            if "(bigAmount)" in paraText:
                paraText = paraText.replace("(bigAmount)", "______")
            if "(borrowPurpose)" in paraText:
                paraText = paraText.replace("(borrowPurpose)", "${dyw}$")
            if "(baseRate)" in paraText:
                paraText = paraText.replace("(baseRate)", "______")
            if "(floatRate)" in paraText:
                paraText = paraText.replace("(floatRate)", "______")
            if "(bankAccountName)" in paraText:
                paraText = paraText.replace("(bankAccountName)", "______")
            if "(bankAccountNo)" in paraText:
                paraText = paraText.replace("(bankAccountNo)", "______")
            if "(bankName)" in paraText:
                paraText = paraText.replace("(bankName)", "______")

        contractHtmlFile.write("<p>" + paraText.strip() + "</p>\n" if paraText.strip() != '' else "<br/>\n")

    contractHtmlFile.write("<br/>\n" +
                           "<p>甲方：${borrowerName}$</p>\n" +
                           "<p>日期：_____</p>\n" +
                           "<br/>\n" +
                           "<p>乙方：（在此签名）</p>\n" +
                           "<p>日期：_____</p>\n" +
                           "<br/>\n" +
                           "<p>丙方：浙江金麦穗互联网金融信息服务有限公司</p>\n" +
                           "<p style=\"position:relative\"><img class=\"imageornot\" src=\"newjms.png\" "
                           "width=\"120\" style=\"position:absolute;left:180px;top:-56px;\"></p>\n" +
                           "<p>法定代表人：曾文杰</p>\n" +
                           "<p>日期：_____</p>")


    # 输出每一段的内容
    for para in sourceFile.paragraphs:
        paraText = para.text

        if "START" in paraText:
            readFile = True
            continue

        if "END" in paraText:
            readFile = False
            continue

        if not readFile:
            continue

        if readFile:
            if "(holder)" in paraText:
                paraText = paraText.replace("(holder)", "${blselfitem.holder!\"_______\"}")
            if "(amount)" in paraText:
                paraText = paraText.replace("(amount)", "______")
            if "(bigAmount)" in paraText:
                paraText = paraText.replace("(bigAmount)", "______")
            if "(borrowPurpose)" in paraText:
                paraText = paraText.replace("(borrowPurpose)", "${(blselfitem.dyw)!\"________\"}")
            if "(baseRate)" in paraText:
                paraText = paraText.replace("(baseRate)", "______")
            if "(floatRate)" in paraText:
                paraText = paraText.replace("(floatRate)", "______")
            if "(bankAccountName)" in paraText:
                paraText = paraText.replace("(bankAccountName)", "<#if '${blselfitem.whitebankaccname!\"\"}' == ''>${(blselfitem.bankaccname)!\"________\"}<#else>${(blselfitem.whitebankaccname)!\"________\"}</#if>")
            if "(bankAccountNo)" in paraText:
                paraText = paraText.replace("(bankAccountNo)", "<#if '${blselfitem.whitebankaccno!\"\"}' == ''>${(blselfitem.bankaccno)!\"________\"}<#else>${(blselfitem.whitebankaccno)!\"________\"}</#if>")
            if "(bankName)" in paraText:
                paraText = paraText.replace("(bankName)", "<#if '${blselfitem.whitebankname!\"\"}' == ''>${(blselfitem.bankname)!\"________\"}<#else>${(blselfitem.whitebankname)!\"________\"}</#if>")

            contractFtlFile.write("<p>" + paraText.strip() + "</p>\n" if paraText.strip() != '' else "<br/>\n")

    contractFtlFile.write("<br/>\n" +
                          "<p>甲方：${(blselfitem.realname)!\"________\"}</p>\n" +
                          "<p>日期：_____</p>\n" +
                          "<br/>\n" +
                          "<p>乙方：（在此签名）</p>\n" +
                          "<p>日期：_____</p>\n" +
                          "<br/>\n" +
                          "<p>丙方：浙江金麦穗互联网金融信息服务有限公司</p>\n" +
                          "<p style=\"position:relative\"><img class=\"imageornot\" src=\"newjms.png\" "
                          "width=\"120\" style=\"position:absolute;left:180px;top:-56px;\"></p>\n" +
                          "<p>法定代表人：曾文杰</p>\n" +
                          "<p>日期：_____</p>")

    contractTxtFile.flush()
    contractHtmlFile.flush()
    contractFtlFile.flush()
    contractTxtFile.close()
    contractHtmlFile.close()
    contractFtlFile.close()

    return True
